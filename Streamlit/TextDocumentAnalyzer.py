"""
Project Name: Text DOcument Analyzer
Module Name and Version: App25 v1.0
Author(s): Srinivas K M
Created On: 09-May-2024
Last Modified On: 13-May-2025
Last Modified by: Srinivas K M
Description: Application for correcting grammar in documents while preserving formatting
"""

import os
import re
import json
import time
import zipfile
import io
import requests
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from collections import Counter
import tempfile
import shutil
from pathlib import Path
import logging
import copy
import xml.etree.ElementTree as ET
from lxml import etree as lxml_etree
import streamlit as st

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'p': 'http://schemas.openxmlformats.org/presentationml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc': 'http://purl.org/dc/elements/1.1/',
    'dcterms': 'http://purl.org/dc/terms/',
    'dcmitype': 'http://purl.org/dc/dcmitype/',
    'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}
for prefix, uri in NAMESPACES.items():
    ET.register_namespace(prefix, uri)

class DocxValidator:
    
    def __init__(self):
        self.namespaces = NAMESPACES
    
    def ValidateAndFixDocx(self, isoInputFile, isoOutputFile=None):
        if isoOutputFile is None:
            isoInputPath = Path(isoInputFile)
            isoOutputFile = str(isoInputPath.parent / f"fixed_{isoInputPath.name}")
        with tempfile.TemporaryDirectory() as isoTempDir:
            isoTempPath = Path(isoTempDir)
            try:
                self._ExtractDocx(isoInputFile, isoTempPath)
            except zipfile.BadZipFile:
                logger.error("Invalid ZIP file structure in DOCX")
                return None
            if not self._ValidateAndFixXmlFiles(isoTempPath):
                logger.error("Failed to validate XML files")
                return None
            if not self._ValidateRelationships(isoTempPath):
                logger.warning("Issues found in relationship files")
            self._CreateDocx(isoTempPath, isoOutputFile)
            return isoOutputFile
    
    def _ExtractDocx(self, isoDocxFile, isoExtractPath):
        with zipfile.ZipFile(isoDocxFile, 'r') as isoZipRef:
            isoZipRef.extractall(isoExtractPath)
    
    def _ValidateAndFixXmlFiles(self, isoExtractPath):
        isoXmlFiles = list(isoExtractPath.glob('**/*.xml'))
        isoXmlFiles.extend(isoExtractPath.glob('**/*.rels'))
        isoSuccess = True
        for isoXmlFile in isoXmlFiles:
            try:
                isoParser = lxml_etree.XMLParser(recover=True)
                isoTree = lxml_etree.parse(str(isoXmlFile), isoParser)
                if len(isoParser.error_log) > 0:
                    logger.warning(f"Fixed {len(isoParser.error_log)} XML errors in {isoXmlFile.name}")
                    isoTree.write(str(isoXmlFile), 
                                 encoding='UTF-8', 
                                 xml_declaration=True,
                                 pretty_print=True)
            except Exception as isoEx:
                logger.error(f"Failed to parse {isoXmlFile.name}: {str(isoEx)}")
                isoSuccess = False
        return isoSuccess
    
    def _ValidateRelationships(self, isoExtractPath):
        isoRelsFiles = list(isoExtractPath.glob('**/*.rels'))
        isoSuccess = True
        for isoRelsFile in isoRelsFiles:
            try:
                isoTree = ET.parse(isoRelsFile)
                isoRoot = isoTree.getroot()
                isoRelationships = isoRoot.findall('.//{*}Relationship')
                isoRelIds = set()
                for isoRel in isoRelationships:
                    isoRelId = isoRel.get('Id')
                    isoTarget = isoRel.get('Target')
                    if isoRelId in isoRelIds:
                        isoNewId = f"R{len(isoRelIds) + 1}"
                        logger.warning(f"Fixed duplicate relationship ID in {isoRelsFile.name}: {isoRelId} -> {isoNewId}")
                        isoRel.set('Id', isoNewId)
                    isoRelIds.add(isoRelId)
                    if isoTarget and not isoTarget.startswith('http') and not isoTarget.startswith('/'):
                        isoTargetPath = isoRelsFile.parent.parent / isoTarget
                        if not isoTargetPath.exists():
                            logger.warning(f"Relationship target not found: {isoTarget}")
                isoTree.write(isoRelsFile, encoding='UTF-8', xml_declaration=True)
            except Exception as isoEx:
                logger.error(f"Failed to process relationships in {isoRelsFile.name}: {str(isoEx)}")
                isoSuccess = False
        return isoSuccess
    
    def _CreateDocx(self, isoSourceDir, isoOutputFile):
        with zipfile.ZipFile(isoOutputFile, 'w', zipfile.ZIP_DEFLATED) as isoZipOut:
            isoContentTypesPath = isoSourceDir / "[Content_Types].xml"
            if isoContentTypesPath.exists():
                isoZipOut.write(isoContentTypesPath, arcname="[Content_Types].xml")
            isoRelsPath = isoSourceDir / "_rels" / ".rels"
            if isoRelsPath.exists():
                isoZipOut.write(isoRelsPath, arcname="_rels/.rels")
            for isoRoot, isoDirs, isoFiles in os.walk(isoSourceDir):
                for isoFile in isoFiles:
                    isoFilePath = Path(isoRoot) / isoFile
                    isoRelPath = isoFilePath.relative_to(isoSourceDir)
                    if str(isoRelPath) == "[Content_Types].xml" or str(isoRelPath) == "_rels/.rels":
                        continue
                    isoZipOut.write(isoFilePath, arcname=str(isoRelPath))


class XMLDocumentCorrector:
    
    def __init__(self, isoCorrector):
        self.corrector = isoCorrector
        self.namespaces = NAMESPACES
    
    def CorrectDocument(self, isoInputFile, isoOutputFile=None):
        with tempfile.TemporaryDirectory() as isoTempDir:
            isoTempPath = Path(isoTempDir)
            if isinstance(isoInputFile, io.BytesIO):
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as isoTempFile:
                    isoTempFile.write(isoInputFile.getvalue())
                    isoTempFilePath = isoTempFile.name
            else:
                isoTempFilePath = isoInputFile
            try:
                with zipfile.ZipFile(isoTempFilePath, 'r') as isoZipRef:
                    isoZipRef.extractall(isoTempPath)
                self._ValidateXmlFiles(isoTempPath)
                self._CorrectDocumentXml(isoTempPath)
                self._CorrectHeadersFooters(isoTempPath)
                self._CorrectOtherParts(isoTempPath)
                self._ValidateXmlFiles(isoTempPath)
                if isoOutputFile is None:
                    isoOutputBuffer = io.BytesIO()
                    self._CreateDocxProperly(isoTempPath, isoOutputBuffer)
                    isoOutputBuffer.seek(0)
                    return isoOutputBuffer
                else:
                    self._CreateDocxProperly(isoTempPath, isoOutputFile)
                    return isoOutputFile
            finally:
                if isinstance(isoInputFile, io.BytesIO) and os.path.exists(isoTempFilePath):
                    try:
                        os.unlink(isoTempFilePath)
                    except:
                        pass
            
    def _ValidateXmlFiles(self, isoExtractPath):
        isoXmlFiles = list(isoExtractPath.glob('**/*.xml'))
        isoXmlFiles.extend(isoExtractPath.glob('**/*.rels'))
        for isoXmlFile in isoXmlFiles:
            try:
                isoParser = lxml_etree.XMLParser(recover=True, remove_blank_text=True)
                isoTree = lxml_etree.parse(str(isoXmlFile), isoParser)
                if len(isoParser.error_log) > 0:
                    logger.warning(f"Fixed {len(isoParser.error_log)} XML errors in {isoXmlFile.name}")
                    isoTree.write(str(isoXmlFile), 
                                 encoding='UTF-8', 
                                 xml_declaration=True,
                                 pretty_print=True)
            except Exception as isoEx:
                logger.error(f"Failed to validate {isoXmlFile.name}: {str(isoEx)}")
            
    def _CorrectDocumentXml(self, isoExtractPath):
        isoDocXmlPath = isoExtractPath / 'word' / 'document.xml'
        
        if not isoDocXmlPath.exists():
            return
        try:
            isoParser = lxml_etree.XMLParser(recover=True, remove_blank_text=True)
            isoTree = lxml_etree.parse(str(isoDocXmlPath), isoParser)
            isoRoot = isoTree.getroot()
            isoParagraphs = isoRoot.xpath('.//w:p', namespaces=self.namespaces)
            isoTables = isoRoot.xpath('.//w:tbl', namespaces=self.namespaces)
            self._ProcessParagraphsInBatches(isoParagraphs, batchSize=10)
            for isoTable in isoTables:
                self._ProcessTableOptimized(isoTable)
            isoTree.write(str(isoDocXmlPath), 
                        encoding='UTF-8', 
                        xml_declaration=True,
                        pretty_print=True)
        except Exception as isoEx:
            logger.exception(f"Error correcting document.xml: {str(isoEx)}")

    def _ProcessParagraphsInBatches(self, isoParagraphs, batchSize=10):
        for i in range(0, len(isoParagraphs), batchSize):
            isoBatch = isoParagraphs[i:i+batchSize]
            isoSectionTexts = []
            isoSectionMappings = []
            for isoParagraph in isoBatch:
                isoTextElements = isoParagraph.xpath('.//w:t', namespaces=self.namespaces)
                if not isoTextElements:
                    continue
                isoParagraphText = ''.join(isoElement.text if isoElement.text else '' for isoElement in isoTextElements)
                if not isoParagraphText.strip():
                    continue
                isoSectionTexts.append(isoParagraphText)
                isoSectionMappings.append({
                    'paragraph': isoParagraph,
                    'elements': isoTextElements,
                    'Original_text': isoParagraphText
                })
            if not isoSectionTexts:
                continue
            isoBatchText = "\n==SECTION_BREAK==\n".join(isoSectionTexts)
            isoAdditionalInstructions = ""
            if hasattr(self.corrector, 'additionalInstructions'):
                isoAdditionalInstructions = self.corrector.additionalInstructions
            isoCorrectedBatch = self.corrector.CorrectGrammar(isoBatchText, isoAdditionalInstructions)
            isoCorrectedSections = isoCorrectedBatch.split("\n==SECTION_BREAK==\n")
            if len(isoCorrectedSections) != len(isoSectionMappings):
                for isoMapping in isoSectionMappings:
                    self._ProcessSingleParagraph(isoMapping['paragraph'])
                continue
            for i, (isoCorrectedSection, isoMapping) in enumerate(zip(isoCorrectedSections, isoSectionMappings)):
                if isoCorrectedSection == isoMapping['Original_text']:
                    continue
                self._DistributeTextToParagraph(isoMapping['elements'], isoMapping['Original_text'], isoCorrectedSection)

    def _DistributeTextToParagraph(self, isoTextElements, isoOriginalText, isoCorrectedText):
        if not isoOriginalText or not isoCorrectedText:
            if isoTextElements:
                isoTextElements[0].text = isoCorrectedText or ""
                for isoElem in isoTextElements[1:]:
                    isoElem.text = ""
            return
        isoElementMap = []
        isoPosition = 0
        for isoElement in isoTextElements:
            isoElemText = isoElement.text or ""
            isoElemLength = len(isoElemText)
            if isoElemLength > 0:
                isoElementMap.append({
                    'element': isoElement,
                    'start_pos': isoPosition,
                    'end_pos': isoPosition + isoElemLength,
                    'Original_text': isoElemText
                })
                isoPosition += isoElemLength
        isoRatio = len(isoCorrectedText) / len(isoOriginalText)
        for isoElemInfo in isoElementMap:
            isoStartRatio = isoElemInfo['start_pos'] / len(isoOriginalText)
            isoEndRatio = isoElemInfo['end_pos'] / len(isoOriginalText)
            isoNewStart = min(int(isoStartRatio * len(isoCorrectedText)), len(isoCorrectedText))
            isoNewEnd = min(int(isoEndRatio * len(isoCorrectedText)), len(isoCorrectedText))
            isoNewText = isoCorrectedText[isoNewStart:isoNewEnd]
            isoElemInfo['element'].text = isoNewText

    def _ProcessTableOptimized(self, isoTable):
        isoCells = isoTable.xpath('.//w:tc', namespaces=self.namespaces)
        for i in range(0, len(isoCells), 5): 
            isoBatch = isoCells[i:i+5]
            isoCellParagraphs = []
            for isoCell in isoBatch:
                isoParagraphs = isoCell.xpath('.//w:p', namespaces=self.namespaces)
                isoCellParagraphs.extend(isoParagraphs)
            if isoCellParagraphs:
                self._ProcessParagraphsInBatches(isoCellParagraphs, batchSize=5)

    def _ProcessSingleParagraph(self, isoParagraph):
        isoTextElements = isoParagraph.xpath('.//w:t', namespaces=self.namespaces)
        if not isoTextElements:
            return
        isoParagraphText = ''.join(isoElement.text if isoElement.text else '' for isoElement in isoTextElements)
        if not isoParagraphText.strip():
            return
        isoAdditionalInstructions = ""
        if hasattr(self.corrector, 'additionalInstructions'):
            isoAdditionalInstructions = self.corrector.additionalInstructions
        isoCorrectedText = self.corrector.CorrectGrammar(isoParagraphText, isoAdditionalInstructions)
        if isoCorrectedText == isoParagraphText:
            return
        self._DistributeTextToParagraph(isoTextElements, isoParagraphText, isoCorrectedText)
    
    def _CorrectHeadersFooters(self, isoExtractPath):
        isoHeaderFooterPaths = list(isoExtractPath.glob('word/header*.xml'))
        isoHeaderFooterPaths.extend(isoExtractPath.glob('word/footer*.xml'))
        for isoPath in isoHeaderFooterPaths:
            try:
                isoParser = lxml_etree.XMLParser(recover=True, remove_blank_text=True)
                isoTree = lxml_etree.parse(str(isoPath), isoParser)
                isoRoot = isoTree.getroot()
                for isoParagraph in isoRoot.xpath('.//w:p', namespaces=self.namespaces):
                    self._ProcessParagraphLxml(isoParagraph)
                isoTree.write(str(isoPath), 
                             encoding='UTF-8', 
                             xml_declaration=True,
                             pretty_print=True)
            except Exception as isoEx:
                logger.error(f"Error correcting header/footer {isoPath.name}: {str(isoEx)}")
    
    def _CorrectOtherParts(self, isoExtractPath):
        isoOtherParts = ['word/comments.xml', 'word/footnotes.xml', 'word/endnotes.xml']
        for isoPartName in isoOtherParts:
            isoPartPath = isoExtractPath / isoPartName
            if not isoPartPath.exists():
                continue
            try:
                isoParser = lxml_etree.XMLParser(recover=True, remove_blank_text=True)
                isoTree = lxml_etree.parse(str(isoPartPath), isoParser)
                isoRoot = isoTree.getroot()
                for isoParagraph in isoRoot.xpath('.//w:p', namespaces=self.namespaces):
                    self._ProcessParagraphLxml(isoParagraph)
                isoTree.write(str(isoPartPath), 
                             encoding='UTF-8', 
                             xml_declaration=True,
                             pretty_print=True)
            except Exception as isoEx:
                logger.error(f"Error correcting {isoPartName}: {str(isoEx)}")
    
    def _ProcessParagraphLxml(self, isoParagraph):
        isoTextElements = isoParagraph.xpath('.//w:t', namespaces=self.namespaces)
        if not isoTextElements:
            return
        isoParagraphText = ''.join(isoElement.text if isoElement.text else '' for isoElement in isoTextElements)
        if not isoParagraphText.strip():
            return
        isoAdditionalInstructions = ""
        if hasattr(self.corrector, 'additionalInstructions'):
            isoAdditionalInstructions = self.corrector.additionalInstructions
        isoCorrectedText = self.corrector.CorrectGrammar(isoParagraphText, isoAdditionalInstructions)
        if isoCorrectedText == isoParagraphText:
            return
        self._DistributeCorrectedTextImproved(isoTextElements, isoParagraphText, isoCorrectedText)
    
    def _DistributeCorrectedTextImproved(self, isoTextElements, isoOriginalText, isoCorrectedText):
        try:
            if len(isoOriginalText) == 0 or len(isoCorrectedText) == 0:
                if isoTextElements:
                    isoTextElements[0].text = isoCorrectedText
                    for isoElem in isoTextElements[1:]:
                        isoElem.text = ""
                return
            isoOriginalLengths = [len(isoElem.text) if isoElem.text else 0 for isoElem in isoTextElements]
            isoTotalOriginalLength = sum(isoOriginalLengths)
            if isoTotalOriginalLength == 0:
                if isoTextElements:
                    isoTextElements[0].text = isoCorrectedText
                return
            isoCharPositions = []
            isoCurrentPos = 0
            for i, isoElement in enumerate(isoTextElements):
                if isoElement.text:
                    for _ in range(len(isoElement.text)):
                        isoCharPositions.append(i)
                        isoCurrentPos += 1
            isoRatio = len(isoCorrectedText) / len(isoOriginalText)
            isoElementTexts = [""] * len(isoTextElements)
            for i, isoChar in enumerate(isoCorrectedText):
                isoOrigPos = min(int(i / isoRatio), len(isoCharPositions) - 1)
                isoElementIndex = isoCharPositions[isoOrigPos]
                isoElementTexts[isoElementIndex] += isoChar
            for i, isoElement in enumerate(isoTextElements):
                isoElement.text = isoElementTexts[i]  
        except Exception as isoEx:
            logger.error(f"Error in text distribution: {str(isoEx)}")
            if isoTextElements:
                isoTextElements[0].text = isoCorrectedText
                for isoElem in isoTextElements[1:]:
                    isoElem.text = ""
    
    def _CreateDocxProperly(self, isoSourceDir, isoOutputFile):
        isoCloseFile = False
        if isinstance(isoOutputFile, str):
            isoOutputFile = open(isoOutputFile, 'wb')
            isoCloseFile = True
        try:
            with zipfile.ZipFile(isoOutputFile, 'w', zipfile.ZIP_DEFLATED) as isoZipOut:
                isoContentTypesPath = isoSourceDir / "[Content_Types].xml"
                if isoContentTypesPath.exists():
                    isoZipOut.write(isoContentTypesPath, arcname="[Content_Types].xml")
                isoRelsPath = isoSourceDir / "_rels" / ".rels"
                if isoRelsPath.exists():
                    isoZipOut.write(isoRelsPath, arcname="_rels/.rels")
                isoDocRelsPath = isoSourceDir / "word" / "_rels" / "document.xml.rels"
                if isoDocRelsPath.exists():
                    isoZipOut.write(isoDocRelsPath, arcname="word/_rels/document.xml.rels")
                isoDocXmlPath = isoSourceDir / "word" / "document.xml"
                if isoDocXmlPath.exists():
                    isoZipOut.write(isoDocXmlPath, arcname="word/document.xml")
                for isoFilePath in (isoSourceDir / "_rels").glob("*"):
                    if isoFilePath.is_file() and isoFilePath.name != ".rels":
                        isoArcname = f"_rels/{isoFilePath.name}"
                        isoZipOut.write(isoFilePath, arcname=isoArcname)
                isoWordRelsDir = isoSourceDir / "word" / "_rels"
                if isoWordRelsDir.exists():
                    for isoFilePath in isoWordRelsDir.glob("*"):
                        if isoFilePath.is_file() and isoFilePath.name != "document.xml.rels":
                            isoArcname = f"word/_rels/{isoFilePath.name}"
                            isoZipOut.write(isoFilePath, arcname=isoArcname)
                for isoFilePath in (isoSourceDir / "word").glob("*"):
                    if isoFilePath.is_file() and isoFilePath.name != "document.xml":
                        isoArcname = f"word/{isoFilePath.name}"
                        isoZipOut.write(isoFilePath, arcname=isoArcname)
                for isoDirPath in (isoSourceDir / "word").glob("*"):
                    if isoDirPath.is_dir() and isoDirPath.name != "_rels":
                        for isoFilePath in isoDirPath.glob("**/*"):
                            if isoFilePath.is_file():
                                isoRelPath = isoFilePath.relative_to(isoSourceDir)
                                isoZipOut.write(isoFilePath, arcname=str(isoRelPath))
                for isoPath in isoSourceDir.glob("*"):
                    if isoPath.name not in ["[Content_Types].xml", "_rels", "word"]:
                        if isoPath.is_dir():
                            for isoFilePath in isoPath.glob("**/*"):
                                if isoFilePath.is_file():
                                    isoRelPath = isoFilePath.relative_to(isoSourceDir)
                                    isoZipOut.write(isoFilePath, arcname=str(isoRelPath))
                        else:
                            isoZipOut.write(isoPath, arcname=isoPath.name)
        finally:
            if isoCloseFile:
                isoOutputFile.close()

class GrammarCorrectorForm:
    
    def __init__(self, isoApiUrl="http://localhost:1234/v1/chat/completions", 
                 isoModel="gemma-3-4b-it-qat", 
                 isoTemperature=0.3,
                 isoMaxTokens=-1):
        self.apiUrl = isoApiUrl
        self.model = isoModel
        self.temperature = isoTemperature
        self.maxTokens = isoMaxTokens
        self._InitializeGrammarRules()
    
    def _InitializeGrammarRules(self):
        self.grammarRuleDefinitions = {
            'punctuation': {
                'title': 'Punctuation Errors',
                'definition': 'Errors in the use of commas, periods, semicolons, and other punctuation marks.',
                'examples': [
                    {'incorrect': 'I went to the store I bought milk.', 'correct': 'I went to the store. I bought milk.'},
                    {'incorrect': 'However she disagreed.', 'correct': 'However, she disagreed.'}
                ]
            },
            'capitalization': {
                'title': 'Capitalization Errors',
                'definition': 'Incorrect use of capital letters for proper nouns, titles, and the start of sentences.',
                'examples': [
                    {'incorrect': 'the capital of france is paris.', 'correct': 'The capital of France is Paris.'},
                    {'incorrect': 'i love reading books.', 'correct': 'I love reading books.'}
                ]
            },
            'verb_tense': {
                'title': 'Verb Tense Errors',
                'definition': 'Inconsistent or incorrect use of verb tenses within a sentence or paragraph.',
                'examples': [
                    {'incorrect': 'Yesterday, I go to the store and bought milk.', 'correct': 'Yesterday, I went to the store and bought milk.'},
                    {'incorrect': 'She has worked here since she starts her career.', 'correct': 'She has worked here since she started her career.'}
                ]
            },
            'subject_verb_agreement': {
                'title': 'Subject-Verb Agreement Errors',
                'definition': 'The subject and verb must agree in number (singular or plural).',
                'examples': [
                    {'incorrect': 'The team are playing well.', 'correct': 'The team is playing well.'},
                    {'incorrect': 'One of the books were missing.', 'correct': 'One of the books was missing.'}
                ]
            },
            'article': {
                'title': 'Article Errors',
                'definition': 'Incorrect use or omission of articles (a, an, the).',
                'examples': [
                    {'incorrect': 'I bought a apple.', 'correct': 'I bought an apple.'},
                    {'incorrect': 'She is best student in class.', 'correct': 'She is the best student in the class.'}
                ]
            },
            'preposition': {
                'title': 'Preposition Errors',
                'definition': 'Incorrect use of prepositions (in, on, at, for, etc.) or preposition combinations.',
                'examples': [
                    {'incorrect': 'She arrived to home.', 'correct': 'She arrived at home.'},
                    {'incorrect': 'I will meet you in Monday.', 'correct': 'I will meet you on Monday.'}
                ]
            },
            'run_on': {
                'title': 'Run-on Sentences',
                'definition': 'Two or more independent clauses joined without appropriate punctuation or conjunctions.',
                'examples': [
                    {'incorrect': 'It was raining I decided to stay home.', 'correct': 'It was raining, so I decided to stay home.'},
                    {'incorrect': 'The report is finished we can submit it now.', 'correct': 'The report is finished. We can submit it now.'}
                ]
            },
            'fragment': {
                'title': 'Sentence Fragments',
                'definition': 'Incomplete sentences that lack a subject, verb, or complete thought.',
                'examples': [
                    {'incorrect': 'Because it was raining.', 'correct': 'I stayed home because it was raining.'},
                    {'incorrect': 'Walking to the store.', 'correct': 'I was walking to the store.'}
                ]
            },
            'double_negative': {
                'title': 'Double Negatives',
                'definition': 'Using two negative words or constructions in the same clause.',
                'examples': [
                    {'incorrect': 'I don\'t have no money.', 'correct': 'I don\'t have any money.'},
                    {'incorrect': 'She didn\'t say nothing.', 'correct': 'She didn\'t say anything.'}
                ]
            },
            'redundancy': {
                'title': 'Redundancy',
                'definition': 'Using unnecessary repetition of words or ideas that express the same meaning.',
                'examples': [
                    {'incorrect': 'The completely full restaurant had no empty tables.', 'correct': 'The restaurant was completely full.'},
                    {'incorrect': 'She nodded her head yes.', 'correct': 'She nodded.'}
                ]
            },
            'its_confusion': {
                'title': 'Its/It\'s Confusion',
                'definition': 'Confusing the possessive "its" with the contraction "it\'s" (it is).',
                'examples': [
                    {'incorrect': 'The dog wagged it\'s tail.', 'correct': 'The dog wagged its tail.'},
                    {'incorrect': 'Its going to rain today.', 'correct': 'It\'s going to rain today.'}
                ]
            },
            'their_confusion': {
                'title': 'Their/There/They\'re Confusion',
                'definition': 'Confusing the possessive "their," the location "there," and the contraction "they\'re" (they are).',
                'examples': [
                    {'incorrect': 'Their going to the store.', 'correct': 'They\'re going to the store.'},
                    {'incorrect': 'I put the book over they\'re.', 'correct': 'I put the book over there.'}
                ]
            },
            'your_confusion': {
                'title': 'Your/You\'re Confusion',
                'definition': 'Confusing the possessive "your" with the contraction "you\'re" (you are).',
                'examples': [
                    {'incorrect': 'Your welcome to join us.', 'correct': 'You\'re welcome to join us.'},
                    {'incorrect': 'Can I borrow you\'re pen?', 'correct': 'Can I borrow your pen?'}
                ]
            },
            'passive_voice': {
                'title': 'Passive Voice',
                'definition': 'Using passive voice when active voice would be clearer and more direct.',
                'examples': [
                    {'incorrect': 'The ball was thrown by John.', 'correct': 'John threw the ball.'},
                    {'incorrect': 'Mistakes were made by the team.', 'correct': 'The team made mistakes.'}
                ]
            },
            'wordiness': {
                'title': 'Wordiness',
                'definition': 'Using more words than necessary to express an idea, making writing unclear or tedious.',
                'examples': [
                    {'incorrect': 'Due to the fact that it was raining, I stayed home.', 'correct': 'Because it was raining, I stayed home.'},
                    {'incorrect': 'In spite of the fact that he was tired, he continued working.', 'correct': 'Although he was tired, he continued working.'}
                ]
            },
            'spelling': {
                'title': 'Spelling Errors',
                'definition': 'Incorrectly spelled words.',
                'examples': [
                    {'incorrect': 'I recieved your message.', 'correct': 'I received your message.'},
                    {'incorrect': 'She is very intellegent.', 'correct': 'She is very intelligent.'}
                ]
            },
            'pronoun_reference': {
                'title': 'Pronoun Reference Errors',
                'definition': 'Unclear or incorrect pronoun references that create ambiguity about the antecedent.',
                'examples': [
                    {'incorrect': 'John told Bob that he was wrong.', 'correct': 'John told Bob, "You are wrong."'},
                    {'incorrect': 'The dog chased the cat until it was tired.', 'correct': 'The dog chased the cat until the dog was tired.'}
                ]
            },
            'modifier_placement': {
                'title': 'Misplaced Modifiers',
                'definition': 'Modifiers (words/phrases that describe) placed too far from what they modify, creating confusion.',
                'examples': [
                    {'incorrect': 'Walking down the street, the trees were beautiful.', 'correct': 'Walking down the street, I thought the trees were beautiful.'},
                    {'incorrect': 'I only ate the vegetables.', 'correct': 'I ate only the vegetables.'}
                ]
            },
            'parallelism': {
                'title': 'Parallelism Errors',
                'definition': 'Lack of parallel structure in pairs or series of related words, phrases, or clauses.',
                'examples': [
                    {'incorrect': 'She likes swimming, running, and to ride bikes.', 'correct': 'She likes swimming, running, and riding bikes.'},
                    {'incorrect': 'The job requires a worker who is skilled, experienced, and has good communication.', 'correct': 'The job requires a worker who is skilled, experienced, and communicative.'}
                ]
            },
            'comma_splice': {
                'title': 'Comma Splice',
                'definition': 'Using only a comma to join two independent clauses without a coordinating conjunction.',
                'examples': [
                    {'incorrect': 'It was raining, I stayed home.', 'correct': 'It was raining, so I stayed home.'},
                    {'incorrect': 'She is intelligent, she works hard.', 'correct': 'She is intelligent, and she works hard.'}
                ]
            },
            'apostrophe': {
                'title': 'Apostrophe Errors',
                'definition': 'Incorrect use of apostrophes for possessives and contractions.',
                'examples': [
                    {'incorrect': 'The dogs collar was red.', 'correct': 'The dog\'s collar was red.'},
                    {'incorrect': 'The childrens toys were everywhere.', 'correct': 'The children\'s toys were everywhere.'}
                ]
            },
            'noun_agreement': {
                'title': 'Noun Agreement Errors',
                'definition': 'Issues with singular/plural forms of nouns, especially with collective nouns or irregular plurals.',
                'examples': [
                    {'incorrect': 'Five person were at the meeting.', 'correct': 'Five people were at the meeting.'},
                    {'incorrect': 'The criterias were not met.', 'correct': 'The criteria were not met.'}
                ]
            },
            'verb_form': {
                'title': 'Incorrect Verb Forms',
                'definition': 'Using the wrong form of a verb, especially with irregular verbs.',
                'examples': [
                    {'incorrect': 'She has went to the store.', 'correct': 'She has gone to the store.'},
                    {'incorrect': 'They teached us everything.', 'correct': 'They taught us everything.'}
                ]
            },
            'word_choice': {
                'title': 'Word Choice Errors',
                'definition': 'Using words incorrectly or choosing words that don\'t convey the intended meaning.',
                'examples': [
                    {'incorrect': 'The effect of his hard work was getting a promotion.', 'correct': 'The result of his hard work was getting a promotion.'},
                    {'incorrect': 'She brought her point across clearly.', 'correct': 'She got her point across clearly.'}
                ]
            },
            'split_infinitive': {
                'title': 'Split Infinitives',
                'definition': 'Placing a word (usually an adverb) between "to" and the verb in an infinitive.',
                'examples': [
                    {'incorrect': 'She wanted to quickly finish the task.', 'correct': 'She wanted to finish the task quickly.'},
                    {'incorrect': 'He decided to carefully review the document.', 'correct': 'He decided to review the document carefully.'}
                ]
            },
            'dangling_participle': {
                'title': 'Dangling Participles',
                'definition': 'A participle or participial phrase that doesn\'t clearly or logically modify a noun or pronoun.',
                'examples': [
                    {'incorrect': 'Running down the street, the bus was missed.', 'correct': 'Running down the street, I missed the bus.'},
                    {'incorrect': 'Having finished the assignment, the TV was turned on.', 'correct': 'Having finished the assignment, she turned on the TV.'}
                ]
            }
        }
        self.isoErrorPatterns = {
            'punctuation': r'[,.!?;:"]',
            'capitalization': r'\b[A-Z][a-z]+\b|^\s*[a-z]',
            'verb_tense': r'\b(is|are|was|were|have|has|had|do|does|did|go|goes|went|seen|saw|eaten|ate)\b',
            'subject_verb_agreement': r'\b(they|we|you)\s+(is|was|has|does)\b|\b(he|she|it)\s+(are|were|have|do)\b',
            'article': r'\b(a|an|the)\b',
            'preposition': r'\b(in|on|at|for|with|by|about|under|over|between|through|after|before|during)\b',
            'run_on': r'[^.!?;:]+[,]\s+[^.!?;:]+[,]\s+[^.!?;:]+[.]\s*',
            'fragment': r'(?<=[.!?])\s+\b(Because|Since|Although|If|When|While|Unless)\b[^.!?]*(?<![.!?])\s',
            'double_negative': r'\b(not|never|no|nobody|nothing|nowhere|neither|nor)\b[^.!?]*\b(not|never|no|nobody|nothing|nowhere|neither|nor)\b',
            'redundancy': r'\b(\w+)\s+\1\b|\b(very\s+really|really\s+very|absolutely\s+essential|completely\s+eliminate)\b',
            'its_confusion': r'\bits\s|it\'s\b',
            'their_confusion': r'\btheir\b|\bthere\b|\bthey\'re\b',
            'your_confusion': r'\byour\b|\byou\'re\b',
            'passive_voice': r'\b(am|is|are|was|were|be|being|been)\s+(\w+ed|written|spoken|done|made|gone|taken|seen|known|given|shown)\b',
            'wordiness': r'\b(at this point in time|due to the fact that|in order to|in the event that|in the process of|on account of the fact that)\b',
            'apostrophe': r'\b\w+s\b|\b\w+\'s\b',
            'comma_splice': r'[^.!?]+,\s+[^.!?]+\.',
            'noun_agreement': r'\b(people|children|men|women|teeth|feet|geese|phenomena|criteria|data)\b',
            'verb_form': r'\b(go|went|gone|do|did|done|see|saw|seen|eat|ate|eaten|write|wrote|written|speak|spoke|spoken)\b',
            'word_choice': r'\b(effect|affect|their|there|they\'re|your|you\'re|its|it\'s|loose|lose|then|than|to|too|two)\b',
            'split_infinitive': r'to\s+\w+ly\s+\w+',
            'dangling_participle': r'\w+ing[^.!?]+,',
            'spelling': r'\b[a-zA-Z]{2,}\b'
        }
    
    def SplitTextIntoSections(self, isoText, isoMaxTokens=1500):
        isoParagraphs = isoText.split('\n')
        isoSections = []
        isoCurrentSection = []
        isoCurrentTokenCount = 0
        for isoPara in isoParagraphs:
            isoParaTokenCount = self.EstimateTokenCount(isoPara)
            if isoParaTokenCount > isoMaxTokens:
                if isoCurrentSection:
                    isoSections.append('\n'.join(isoCurrentSection))
                    isoCurrentSection = []
                    isoCurrentTokenCount = 0
                isoSentences = re.split(r'(?<=[.!?])\s+', isoPara)
                isoSentenceSection = []
                isoSentenceTokenCount = 0
                for isoSentence in isoSentences:
                    isoSentenceTokens = self.EstimateTokenCount(isoSentence)
                    if isoSentenceTokenCount + isoSentenceTokens > isoMaxTokens:
                        if isoSentenceSection:
                            isoSections.append(' '.join(isoSentenceSection))
                            isoSentenceSection = [isoSentence]
                            isoSentenceTokenCount = isoSentenceTokens
                        else:
                            isoSections.append(isoSentence)
                    else:
                        isoSentenceSection.append(isoSentence)
                        isoSentenceTokenCount += isoSentenceTokens
                if isoSentenceSection:
                    isoSections.append(' '.join(isoSentenceSection))
            elif isoCurrentTokenCount + isoParaTokenCount <= isoMaxTokens:
                isoCurrentSection.append(isoPara)
                isoCurrentTokenCount += isoParaTokenCount
            else:
                isoSections.append('\n'.join(isoCurrentSection))
                isoCurrentSection = [isoPara]
                isoCurrentTokenCount = isoParaTokenCount
        if isoCurrentSection:
            isoSections.append('\n'.join(isoCurrentSection))
        return isoSections
    
    def _CreatePrompt(self, isoText, isoAdditionalInstructions=""):
        isoSystemPrompt = """You are an expert grammar correction system. Your task is to correct grammar errors in the provided text while preserving its Original meaning, style, and formatting. 
        Follow these rules:
        1. Fix grammar, punctuation, spelling, and syntax errors only
        2. Do NOT change the meaning or style of the text
        3. Do NOT add or remove content
        4. Do NOT change formatting like bullet points, numbering, or paragraph breaks
        5. Return ONLY the Corrected text without any explanations, rationale, or notes
        6. If the text has no grammatical errors, return it unchanged
        For each correction, ensure the fix is necessary and appropriate for formal writing."""
        if isoAdditionalInstructions:
            isoSystemPrompt += f"\n\nAdditional instructions: {isoAdditionalInstructions}"
        isoMessages = [
            {"role": "system", "content": isoSystemPrompt},
            {"role": "user", "content": f"Correct any grammar errors in the following text:\n\n{isoText}"}
        ]
        return isoMessages
    
    def CorrectGrammar(self, isoText, isoAdditionalInstructions="", isoRetryCount=3):
        if not isoText or isoText.isspace():
            return isoText
        isoMessages = self._CreatePrompt(isoText, isoAdditionalInstructions)
        isoPayload = {
            "model": self.model,
            "messages": isoMessages,
            "temperature": self.temperature,
            "max_tokens": self.maxTokens,
            "stream": False
        }
        for isoAttempt in range(isoRetryCount):
            try:
                isoResponse = requests.post(self.apiUrl, json=isoPayload)
                isoResponse.raise_for_status()
                isoResult = isoResponse.json()
                isoCorrectedText = isoResult["choices"][0]["message"]["content"]
                return isoCorrectedText
            except (requests.RequestException, json.JSONDecodeError, KeyError, IndexError) as isoEx:
                logger.error(f"API error (attempt {isoAttempt+1}/{isoRetryCount}): {str(isoEx)}")
                if isoAttempt < isoRetryCount - 1:
                    time.sleep(2)
                else:
                    logger.error(f"Failed to correct text after {isoRetryCount} attempts. Returning Original text.")
                    return isoText
    
    def CorrectTextInSections(self, isoText, isoAdditionalInstructions="", isoProgressCallback=None):
        isoSections = self.SplitTextIntoSections(isoText)
        isoTotalSections = len(isoSections)
        isoCorrectedSections = []
        for i, isoSection in enumerate(isoSections):
            if isoProgressCallback:
                isoProgressCallback(i / isoTotalSections, f"Correcting section {i+1}/{isoTotalSections}")
            isoCorrectedSection = self.CorrectGrammar(isoSection, isoAdditionalInstructions)
            isoCorrectedSections.append(isoCorrectedSection)
            time.sleep(0.5)
        isoCorrectedText = '\n'.join(isoCorrectedSections)
        return isoCorrectedText

class SafeDocxCreatorForm:

    def CreateSafeDocx(self, isoOriginalText, isoCorrectedText):
        isoDoc = Document()
        isoParagraphs = isoCorrectedText.split('\n')
        for isoParaText in isoParagraphs:
            if isoParaText.strip():  
                isoDoc.add_paragraph(isoParaText)
        isoDocxIo = io.BytesIO()
        isoDoc.save(isoDocxIo)
        isoDocxIo.seek(0)
        return isoDocxIo
    
    def CreateSafeDocxWithFormatting(self, isoOriginalFile, isoCorrectedText):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as isoTempFile:
                isoTempFile.write(isoOriginalFile.read())
                isoOriginalFile.seek(0) 
                isoTempFilePath = isoTempFile.name
            isoOriginalDoc = Document(isoTempFilePath)
            isoNewDoc = Document()
            try:
                isoNewDoc.styles = isoOriginalDoc.styles
                isoNewDoc.core_properties.author = isoOriginalDoc.core_properties.author
                isoNewDoc.core_properties.title = isoOriginalDoc.core_properties.title
                isoNewDoc.core_properties.comments = "Grammar Corrected version"
            except:
                logger.warning("Could not copy document styles")
            isoParagraphs = isoCorrectedText.split('\n')
            isoOriginalParagraphs = [p.text for p in isoOriginalDoc.paragraphs]
            for i, isoParaText in enumerate(isoParagraphs):
                if not isoParaText.strip(): 
                    continue
                isoNewPara = isoNewDoc.add_paragraph(isoParaText)
                try:
                    if i < len(isoOriginalParagraphs):
                        isoOrigPara = isoOriginalDoc.paragraphs[i]
                        if isoOrigPara.style:
                            isoNewPara.style = isoOrigPara.style
                except:
                    pass
            isoDocxIo = io.BytesIO()
            isoNewDoc.save(isoDocxIo)
            isoDocxIo.seek(0)
            os.unlink(isoTempFilePath)
            return isoDocxIo
        except Exception as isoEx:
            logger.error(f"Error creating formatted DOCX: {str(isoEx)}")
            return self.CreateSafeDocx(isoCorrectedText, isoCorrectedText)


class DocumentAnalyzerForm:
    
    def AnalyzeCorrections(self, isoOriginalText, isoCorrectedText, isoCorrector=None):
        if isoOriginalText is None:
            isoOriginalText = ""
        if isoCorrectedText is None:
            isoCorrectedText = ""
        isoOrigParas = [p.strip() for p in isoOriginalText.split('\n') if p.strip()]
        isoCorrParas = [p.strip() for p in isoCorrectedText.split('\n') if p.strip()]
        isoChanges = []
        isoMinLen = min(len(isoOrigParas), len(isoCorrParas))
        for i in range(isoMinLen):
            isoOrig = isoOrigParas[i]
            isoCorr = isoCorrParas[i]
            if isoOrig != isoCorr:
                isoChanges.append({
                    'paragraph': i + 1,
                    'Original': isoOrig,
                    'Corrected': isoCorr,
                    'Reasoning': None,  
                    'DetectedErrorTypes': []  
                })
        for i in range(isoMinLen, len(isoOrigParas)):
            isoChanges.append({
                'paragraph': i + 1,
                'Original': isoOrigParas[i],
                'Corrected': "",  
                'Reasoning': "Added or removed content",
                'DetectedErrorTypes': ["Content Change"]
            })
        for i in range(isoMinLen, len(isoCorrParas)):
            isoChanges.append({
                'paragraph': len(isoOrigParas) + i - isoMinLen + 1,
                'Original': "", 
                'Corrected': isoCorrParas[i],
                'Reasoning': "Added or removed content",
                'DetectedErrorTypes': ["Content Change"]
            })
        if isoCorrector and isoChanges:
            for isoChange in isoChanges:
                if not isoChange['Original'] or not isoChange['Corrected']:
                    continue
                self._DetectParagraphErrorTypes(isoChange)
                try:
                    isoChange['Reasoning'] = self._GenerateReasoningFromDetectedTypes(isoCorrector, isoChange)
                    time.sleep(0.2)
                except Exception as isoEx:
                    logger.error(f"Error generating reasoning: {str(isoEx)}")
                    if isoChange['DetectedErrorTypes']:
                        primaryType = isoChange['DetectedErrorTypes'][0]
                        isoChange['Reasoning'] = f"{primaryType}: Error in {primaryType.lower()}."
                    else:
                        isoChange['Reasoning'] = "Grammar or spelling error"
        isoErrorTypeCounts = Counter()
        for isoChange in isoChanges:
            for isoErrorType in isoChange['DetectedErrorTypes']:
                isoErrorTypeCounts[isoErrorType] += 1
        return {'changes': isoChanges, 'error_types': dict(isoErrorTypeCounts)}

    def _DetectParagraphErrorTypes(self, isoChange):
        isoOrigText = isoChange.get('Original', '')
        isoCorrText = isoChange.get('Corrected', '')
        if not isoOrigText or not isoCorrText:
            return
        isoErrorPatterns = self._GetErrorPatterns()
        isoDetectedTypes = []
        for isoErrorType, isoPattern in isoErrorPatterns.items():
            try:
                isoOrigMatches = set(re.findall(isoPattern, isoOrigText))
                isoCorrMatches = set(re.findall(isoPattern, isoCorrText))
                if isoOrigMatches != isoCorrMatches:
                    isoDetectedTypes.append(isoErrorType)
            except Exception as isoEx:
                logger.error(f"Error checking pattern {isoErrorType}: {str(isoEx)}")
                continue
        isoChange['DetectedErrorTypes'] = isoDetectedTypes

    def _GenerateReasoningFromDetectedTypes(self, isoCorrector, isoChange):
        isoDetectedTypes = isoChange.get('DetectedErrorTypes', [])
        if not isoDetectedTypes:
            return self._GenerateErrorReasoningWithLM(isoCorrector, isoChange)
        primaryType = isoDetectedTypes[0]
        if len(isoDetectedTypes) == 1:
            return f"{primaryType}: The primary error is {primaryType.lower()}."
        else:
            errorList = ", ".join(isoDetectedTypes[:-1]) + " and " + isoDetectedTypes[-1]
            return f"{primaryType}: {errorList} errors were corrected."

    def _GenerateErrorReasoningWithLM(self, isoCorrector, isoChange):
        try:
            isoPrompt = f"""Identify specific grammar and spelling errors corrected in the text change. 
            Answer in a SINGLE brief sentence identifying the primary error type.
            Original text: "{isoChange['Original']}"
            Corrected text: "{isoChange['Corrected']}"
            Error type:"""
            isoMessages = [
                {"role": "system", "content": "You are a grammar teacher providing very concise error identification in one brief sentence. Identify the specific error from these categories: Punctuation, Capitalization, Verb Tense, Subject-Verb Agreement, Article Usage, Preposition Usage, Run-on Sentences, Sentence Fragments, Double Negatives, Redundancy, Pronoun Confusion (its/it's, their/there/they're, your/you're), Passive Voice, Wordiness, Apostrophe Errors, Comma Splices, Noun Agreement, Verb Form, Word Choice, Split Infinitives, Dangling Participles, or Spelling."},
                {"role": "user", "content": isoPrompt}
            ]
            isoPayload = {
                "model": isoCorrector.model,
                "messages": isoMessages,
                "temperature": 0.2,  
                "max_tokens": 50,   
                "stream": False
            }
            isoResponse = requests.post(isoCorrector.apiUrl, json=isoPayload)
            isoResponse.raise_for_status()
            isoResult = isoResponse.json()
            isoExplanation = isoResult["choices"][0]["message"]["content"].strip()
            for isoErrorType in self._GetErrorPatterns().keys():
                if isoErrorType.lower() in isoExplanation.lower():
                    if isoErrorType not in isoChange['DetectedErrorTypes']:
                        isoChange['DetectedErrorTypes'].append(isoErrorType)
            if len(isoExplanation) > 100:
                isoExplanation = isoExplanation.split('.')[0] + '.'  
            return isoExplanation
        except Exception as isoEx:
            logger.error(f"Error generating error reasoning with LM: {str(isoEx)}")
            return "Grammar or spelling error"
    
    def _GetErrorPatterns(self):
        return {
            'Punctuation': r'[,.!?;:"]',
            'Capitalization': r'\b[A-Z][a-z]+\b|^\s*[a-z]',
            'Verb Tense': r'\b(is|are|was|were|have|has|had|do|does|did|go|goes|went|seen|saw|eaten|ate)\b',
            'Subject Verb Agreement': r'\b(they|we|you)\s+(is|was|has|does)\b|\b(he|she|it)\s+(are|were|have|do)\b',
            'Article': r'\b(a|an|the)\b',
            'Preposition': r'\b(in|on|at|for|with|by|about|under|over|between|through|after|before|during)\b',
            'Run on': r'[^.!?;:]+[,]\s+[^.!?;:]+[,]\s+[^.!?;:]+[.]\s*',
            'Fragment': r'(?<=[.!?])\s+\b(Because|Since|Although|If|When|While|Unless)\b[^.!?]*(?<![.!?])\s',
            'Double Negative': r'\b(not|never|no|nobody|nothing|nowhere|neither|nor)\b[^.!?]*\b(not|never|no|nobody|nothing|nowhere|neither|nor)\b',
            'Redundancy': r'\b(\w+)\s+\1\b|\b(very\s+really|really\s+very|absolutely\s+essential|completely\s+eliminate)\b',
            'Its Confusion': r'\bits\s|it\'s\b',
            'Their Confusion': r'\btheir\b|\bthere\b|\bthey\'re\b',
            'Your Confusion': r'\byour\b|\byou\'re\b',
            'Passive Voice': r'\b(am|is|are|was|were|be|being|been)\s+(\w+ed|written|spoken|done|made|gone|taken|seen|known|given|shown)\b',
            'Wordiness': r'\b(at this point in time|due to the fact that|in order to|in the event that|in the process of|on account of the fact that)\b',
            'Apostrophe': r'\b\w+s\b|\b\w+\'s\b',
            'Comma Splice': r'[^.!?]+,\s+[^.!?]+\.',
            'Noun Agreement': r'\b(people|children|men|women|teeth|feet|geese|phenomena|criteria|data)\b',
            'Verb Form': r'\b(go|went|gone|do|did|done|see|saw|seen|eat|ate|eaten|write|wrote|written|speak|spoke|spoken)\b',
            'Word Choice': r'\b(effect|affect|their|there|they\'re|your|you\'re|its|it\'s|loose|lose|then|than|to|too|two)\b',
            'Split Infinitive': r'to\s+\w+ly\s+\w+',
            'Dangling Participle': r'\w+ing[^.!?]+,',
            'Spelling': r'\b[a-zA-Z]{2,}\b'
        }
    
    def GenerateSummaryReport(self, isoAnalysis):
        isoTotalChanges = len(isoAnalysis['changes'])
        isoErrorTypes = isoAnalysis['error_types']
        isoSummary = f"Grammar Correction Report\n"
        isoSummary += f"========================\n\n"
        isoSummary += f"Total corrections made: {isoTotalChanges}\n\n"
        if isoErrorTypes:
            isoSummary += f"Error types found:\n"
            isoSortedErrors = sorted(isoErrorTypes.items(), key=lambda x: x[1], reverse=True)
            for isoErrorType, isoCount in isoSortedErrors:
                isoSummary += f"- {isoErrorType}: {isoCount}\n"
        else:
            isoSummary += "No grammar errors were found in the document.\n"
        return isoSummary


class DocumentProcessorForm:

    def __init__(self, isoApiUrl="http://localhost:1234/v1/chat/completions", 
                 isoModel="gemma-3-4b-it-qat", 
                 isoTemperature=0.3):
        self.corrector = GrammarCorrectorForm(isoApiUrl, isoModel, isoTemperature)
        self.validator = DocxValidator()
        self.xmlCorrector = XMLDocumentCorrector(self.corrector)
        self.docxCreator = SafeDocxCreatorForm()
        self.analyzer = DocumentAnalyzerForm()
    
    def ExtractTextFromDoc(self, isoFileData, isoFilename):
        if isoFilename.lower().endswith('.docx'):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as isoTempFile:
                    isoTempFile.write(isoFileData.getvalue())
                    isoTempFilePath = isoTempFile.name
                isoDoc = Document(isoTempFilePath)
                isoText = '\n'.join([p.text for p in isoDoc.paragraphs if p.text.strip()])
                os.unlink(isoTempFilePath)
                return isoText
            except Exception as isoEx:
                logger.error(f"Error extracting text from DOCX: {str(isoEx)}")
                raise ValueError("Could not read the DOCX file. It may be corrupted.")
        elif isoFilename.lower().endswith('.txt'):
            return isoFileData.getvalue().decode('utf-8', errors='ignore')
        elif isoFilename.lower().endswith('.doc'):
            return isoFileData.getvalue().decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {isoFilename}")
    
    def ProcessDocument(self, isoUploadedFile, isoCompatibilityMode="Safe Mode", isoProgressCallback=None, isoAdditionalInstructions=""):
        isoResults = {
            'Original_text': None,
            'Corrected_text': None,
            'Corrected_docx': None,
            'analysis': None,
            'summary': None
        }   
        try:
            if isoProgressCallback:
                isoProgressCallback(0.1, "Extracting text from document...")
            isoIsDocx = isoUploadedFile.name.lower().endswith('.docx')
            isoOriginalText = self.ExtractTextFromDoc(isoUploadedFile, isoUploadedFile.name)
            isoResults['Original_text'] = isoOriginalText
            isoResults['Corrected_text'] = isoOriginalText
            if isoProgressCallback:
                isoProgressCallback(0.2, "Starting grammar correction...")
            if isoIsDocx and isoCompatibilityMode == "Preserve All":
                isoUploadedFile.seek(0)
                isoSectionProgress = lambda p, m: isoProgressCallback(0.2 + p * 0.6, m) if isoProgressCallback else None
                if isoProgressCallback:
                    isoProgressCallback(0.3, "Correcting document while preserving all formatting...")
                isoFileData = io.BytesIO(isoUploadedFile.getvalue())
                self.xmlCorrector.corrector.additionalInstructions = isoAdditionalInstructions
                isoDocxIo = self.xmlCorrector.CorrectDocument(isoFileData)
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as isoTempFile:
                        isoTempFile.write(isoDocxIo.getvalue())
                        isoDocxIo.seek(0)
                        isoTempFilePath = isoTempFile.name
                    isoDoc = Document(isoTempFilePath)
                    isoCorrectedText = '\n'.join([p.text for p in isoDoc.paragraphs if p.text.strip()])
                    os.unlink(isoTempFilePath)
                    if not isoCorrectedText.strip():
                        isoCorrectedText = isoOriginalText
                    isoResults['Corrected_text'] = isoCorrectedText
                except Exception as isoEx:
                    logger.error(f"Error extracting Corrected text: {str(isoEx)}")
                isoResults['Corrected_docx'] = isoDocxIo.getvalue()  
            else:
                isoSectionProgress = lambda p, m: isoProgressCallback(0.2 + p * 0.6, m) if isoProgressCallback else None
                isoCorrectedText = self.corrector.CorrectTextInSections(
                    isoOriginalText, 
                    isoAdditionalInstructions,
                    isoSectionProgress
                )
                if isoCorrectedText is None or not isoCorrectedText.strip():
                    isoCorrectedText = isoOriginalText  
                isoResults['Corrected_text'] = isoCorrectedText
                if isoProgressCallback:
                    isoProgressCallback(0.8, "Creating compatible document...")
                isoUploadedFile.seek(0)
                if isoCompatibilityMode == "Safe Mode":
                    isoDocxIo = self.docxCreator.CreateSafeDocxWithFormatting(isoUploadedFile, isoCorrectedText)
                else: 
                    isoDocxIo = self.docxCreator.CreateSafeDocx(isoOriginalText, isoCorrectedText)
                isoResults['Corrected_docx'] = isoDocxIo.getvalue()
            if isoProgressCallback:
                isoProgressCallback(0.9, "Analyzing corrections...")
            isoAnalysis = self.analyzer.AnalyzeCorrections(
                isoResults['Original_text'] or "", 
                isoResults['Corrected_text'] or "",
                self.corrector  
            )
            isoResults['analysis'] = isoAnalysis
            isoSummary = self.analyzer.GenerateSummaryReport(isoAnalysis)
            isoResults['summary'] = isoSummary
            if isoProgressCallback:
                isoProgressCallback(1.0, "Document processing complete!")
            return isoResults
        except Exception as isoEx:
            logger.exception(f"Error in document processing: {str(isoEx)}")
            raise


class DocumentCorrectionAppView:

    def __init__(self):
        self.SetupPage()
        self.InitSessionState()
        self.processor = None  
    
    def SetupPage(self):
        st.set_page_config(
            page_title="Text Document Analyzer", 
            page_icon="📝", 
            layout="wide"
        )
    
    def InitSessionState(self):
        if 'current_state' not in st.session_state:
            st.session_state.current_state = "upload" 
        if 'processing_results' not in st.session_state:
            st.session_state.processing_results = None
        if 'download_count' not in st.session_state:
            st.session_state.download_count = 0
        if 'filename' not in st.session_state:
            st.session_state.filename = None
        if 'processor' not in st.session_state:
            st.session_state.processor = None
    
    def RenderSidebar(self):
        with st.sidebar:
            st.header("Configuration")
            isoApiUrl = st.text_input("LM Studio API URL", value="http://localhost:1234/v1/chat/completions")
            isoModel = st.text_input("Model Name", value="gemma-3-4b-it-qat")
            isoTemperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.3, step=0.1)
            st.header("Processing Options")
            isoShowComparison = st.checkbox("Show text comparison", value=True)
            isoShowAnalysis = st.checkbox("Show error analysis", value=True)
            st.header("Document Compatibility Mode")
            isoCompatibilityMode = st.radio(
                "Select mode for highest compatibility",
                ["Preserve All", "Safe Mode", "Ultra Safe Mode"],
                index=0,
                help=("Preserve All tries to maintain tables, images, and formatting. "
                      "Safe Mode tries to balance formatting and compatibility. "
                      "Ultra Safe mode creates a completely new document with perfect compatibility but minimal formatting.")
            )
            return {
                'api_url': isoApiUrl,
                'model': isoModel,
                'temperature': isoTemperature,
                'show_comparison': isoShowComparison,
                'show_analysis': isoShowAnalysis,
                'compatibility_mode': isoCompatibilityMode
            }
            
    def RenderUploadPage(self, isoConfig):
        st.title("Text Document Analyzer")
        st.subheader("Correct grammar while preserving document formatting")
        st.markdown("""
        <div style="background-color:#e8f4fd; padding:20px; border-radius:10px; margin-top:25px; margin-bottom:30px; border-left:5px solid #3498db; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">
            <h3 style="color:#3498db; margin-top:0;">Additional Instructions</h3>
            <p style="font-size:16px;">Provide specific instructions for how you want the grammar to be corrected.</p>
            <p style="font-size:14px; color:#666; font-style:italic;">For example: Focus on academic style, avoid contractions, maintain technical terminology, etc.</p>
        </div>
        """, unsafe_allow_html=True)
        isoAdditionalInstructions = st.text_area(
            "Enter your instructions below:",
            placeholder="Example: Focus on academic style, avoid contractions, maintain technical terminology, etc.",
            help="These instructions will be added to the prompt sent to the language model.",
            key="instruction_input_widget"
        )
        st.markdown("<hr style='margin: 30px 0; border-top: 1px solid #eee;'>", unsafe_allow_html=True)
        isoUploadedFile = st.file_uploader(
            "Upload a document (.docx, .doc, .txt)", 
            type=["docx", "doc", "txt"],
            help="Drag and drop your file here or click to browse"
        )
        if isoUploadedFile is not None:
            st.success(f"File uploaded: {isoUploadedFile.name}")
            st.markdown("""
            <style>
            div.stButton > button {
                background-color: #4CAF50;
                color: white;
                padding: 12px 24px;
                border-radius: 8px;
                font-weight: bold;
                transition: all 0.3s;
                display: inline-flex;
                align-items: center;
                justify-content: center;
            }
            div.stButton > button:hover {
                background-color: #45a049;
                box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
            }
            </style>
            """, unsafe_allow_html=True)
            if st.button("Correct Grammar"):
                isoProcessor = DocumentProcessorForm(
                    isoApiUrl=isoConfig['api_url'],
                    isoModel=isoConfig['model'],
                    isoTemperature=isoConfig['temperature']
                )
                st.session_state.processor = isoProcessor
                st.session_state.current_state = "processing"
                st.session_state.filename = isoUploadedFile.name
                st.session_state.user_instructions = isoAdditionalInstructions
                st.session_state.uploaded_file = isoUploadedFile
                st.rerun()
                
    def RenderProcessingPage(self, isoConfig):
        st.title("Processing Document")
        if st.session_state.processor is None:
            st.session_state.processor = DocumentProcessorForm(
                isoApiUrl=isoConfig['api_url'],
                isoModel=isoConfig['model'],
                isoTemperature=isoConfig['temperature']
            )
        st.markdown("""
        <style>
        .stProgress > div > div > div > div {
            background-color: #3498db;
        }
        </style>
        """, unsafe_allow_html=True)
        isoProgressBar = st.progress(0)
        isoStatusText = st.empty()
        if hasattr(st.session_state, 'user_instructions') and st.session_state.user_instructions:
            st.markdown(f"""
            <div style="background-color:#f0f9ff; padding:15px; border-radius:8px; margin-top:20px; margin-bottom:20px; border-left:5px solid #3498db;">
            <h4 style="color:#3498db;">Using Additional Instructions</h4>
            <p style="font-style:italic;">{st.session_state.user_instructions}</p>
            </div>
            """, unsafe_allow_html=True)
        def UpdateProgress(isoProgress, isoMessage):
            isoProgressBar.progress(isoProgress)
            isoStatusText.markdown(f"""
            <div style="background-color:#f8f8f8; padding:10px; border-radius:5px; margin-top:10px;">
            <p style="font-weight:bold; color:#3498db; margin:0;">{isoMessage}</p>
            </div>
            """, unsafe_allow_html=True)
        try:
            isoAdditionalInstructions = ""
            if hasattr(st.session_state, 'user_instructions'):
                isoAdditionalInstructions = st.session_state.user_instructions
            isoResults = st.session_state.processor.ProcessDocument(
                st.session_state.uploaded_file,
                isoConfig['compatibility_mode'],
                UpdateProgress,
                isoAdditionalInstructions
            )
            st.session_state.processing_results = isoResults
            st.session_state.current_state = "results"
            st.session_state.download_count += 1
            st.rerun()
            
        except Exception as isoEx:
            st.error(f"Error processing document: {str(isoEx)}")
            logger.exception("Document processing error")
            st.session_state.current_state = "upload"
            st.rerun()
    
    def RenderResultsPage(self, isoConfig):
        isoResults = st.session_state.processing_results
        st.title("Text Document Analyzer")
        st.success("Grammar correction completed!")
        isoTabs = st.tabs(["Corrected Document", "Text Comparison", "Error Analysis", "Detailed Report"])
        # Tab 1: Download Corrected Document
        with isoTabs[0]:
            st.markdown("### Download Your Corrected Document")
            isoDownloadKey = f"download_{st.session_state.download_count}"
            isoCol1, isoCol2 = st.columns([3, 1])
            with isoCol1:
                st.markdown("""
                <style>
                div.stDownloadButton > button {
                    background-color: #4CAF50;
                    color: white;
                    padding: 10px 24px;
                    border-radius: 8px;
                    font-weight: bold;
                    transition: all 0.3s;
                }
                div.stDownloadButton > button:hover {
                    background-color: #45a049;
                    box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
                }
                </style>
                """, unsafe_allow_html=True)
                st.download_button(
                    label="📄 Download Corrected DOCX Document",
                    data=isoResults['Corrected_docx'],
                    file_name=f"Corrected_{st.session_state.filename}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=isoDownloadKey,
                    help="Compatible Corrected document that should open in all versions of Word"
                )
                isoMarkdownReport = self._GenerateMarkdownReport(isoResults)
                isoBaseFilename = st.session_state.filename
                if '.' in isoBaseFilename:
                    isoBaseFilename = isoBaseFilename.split('.')[0]
                st.download_button(
                    label="📊 Download Error Analysis Report (Markdown)",
                    data=isoMarkdownReport,
                    file_name=f"ErrorReportAnalysis_{isoBaseFilename}.md",
                    mime="text/markdown",
                    key=f"md_download_{st.session_state.download_count}",
                    help="Download a detailed analysis of grammar corrections in Markdown format"
                )
                if isoConfig['compatibility_mode'] == "Preserve All":
                    st.info("Note: The document has been Corrected while preserving all tables, images and formatting.")
                else:
                    st.info("Note: Downloaded files can be safely opened in Microsoft Word. If you have any issues, try using the Preserve All Mode from the sidebar before processing.")
            with isoCol2:
                st.markdown("""
                <style>
                div.stButton > button {
                    background-color: #3498db;
                    color: white;
                    padding: 10px 24px;
                    border-radius: 8px;
                    font-weight: bold;
                    transition: all 0.3s;
                }
                div.stButton > button:hover {
                    background-color: #2980b9;
                    box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
                }
                </style>
                """, unsafe_allow_html=True)
                if st.button("Process New Document"):
                    st.session_state.current_state = "upload"
                    st.session_state.processing_results = None
                    st.rerun()
        # Tab 2: Text comparison if configured and available
        with isoTabs[1]:
            if isoConfig['show_comparison']:
                st.subheader("Text Comparison")
                isoCol1, isoCol2 = st.columns(2)
                with isoCol1:
                    st.markdown("**Original Text**")
                    st.text_area("", value=isoResults['Original_text'], height=300, key="orig_text_area")
                with isoCol2:
                    st.markdown("**Corrected Text**")
                    st.text_area("", value=isoResults['Corrected_text'], height=300, key="corr_text_area")
        # Tab 3: Error analysis if configured and available
        with isoTabs[2]:
            if isoConfig['show_analysis'] and isoResults['analysis']:
                st.subheader("Error Analysis")
                if isoResults['analysis']['error_types']:
                    st.markdown("**Error Types Found**")
                    isoErrorDf = pd.DataFrame({
                        'Error Type': list(isoResults['analysis']['error_types'].keys()),
                        'Count': list(isoResults['analysis']['error_types'].values())
                    })
                    isoErrorDf = isoErrorDf.sort_values('Count', ascending=False)
                    import plotly.express as px
                    isoFig = px.pie(
                        isoErrorDf, 
                        values='Count', 
                        names='Error Type',
                        title='Grammar Error Types',
                        hole=0.4, 
                        color_discrete_sequence=px.colors.qualitative.Plotly,
                    )
                    isoFig.update_layout(
                        title_font_size=18,
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=-0.2,
                            xanchor="center",
                            x=0.5
                        )
                    )
                    isoFig.update_traces(
                        textinfo='percent+label',
                        hoverinfo='label+value',
                        textfont_size=12
                    )
                    st.plotly_chart(isoFig, use_container_width=True)
                    isoTopErrors = isoErrorDf.head(10)  
                    isoFig2 = px.bar(
                        isoTopErrors,
                        y='Error Type',
                        x='Count',
                        title='Top Grammar Error Types',
                        orientation='h',
                        color='Count',
                        color_continuous_scale='Blues',
                    )
                    isoFig2.update_layout(
                        title_font_size=18,
                        yaxis_title="",
                        xaxis_title="Count",
                    )
                    isoFig2.update_traces(
                        hovertemplate='<b>%{y}</b><br>Count: %{x}<extra></extra>',
                        texttemplate='%{x}',
                        textposition='outside'
                    )
                    st.plotly_chart(isoFig2, use_container_width=True)
                isoTotalCorrections = len(isoResults['analysis']['changes'])
                st.markdown(f"""
                <div style="background-color:#f0f9ff; padding:20px; border-radius:10px; border-left:5px solid #3498db;">
                <h3 style="color:#3498db;">Summary Statistics</h3>
                <p style="font-size:18px;">Total corrections made: <span style="font-weight:bold; color:#3498db;">{isoTotalCorrections}</span></p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.info("No grammar errors were found in the document.")
        # Tab 4: Detailed Report
        with isoTabs[3]:
            st.subheader("Detailed Correction Report")
            if isoResults['analysis']['changes']:
                for i, isoChange in enumerate(isoResults['analysis']['changes']):
                    st.markdown(f"""
                    <div style="background-color:#f8f8f8; padding:10px; border-radius:5px; margin-bottom:10px; border-left:3px solid #2ecc71;">
                    <h4>Paragraph {isoChange['paragraph']}</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    isoCol1, isoCol2 = st.columns(2)
                    with isoCol1:
                        st.markdown("""
                        <div style="background-color:#fff4f4; padding:10px; border-radius:5px; border:1px solid #ffcccb;">
                        <p style="font-weight:bold; color:#e74c3c;">Original:</p>
                        </div>
                        """, unsafe_allow_html=True)
                        st.markdown(f"<div style='background-color:#fff; padding:10px; border-radius:5px; border:1px solid #eee;'>{isoChange['Original']}</div>", unsafe_allow_html=True)
                    with isoCol2:
                        st.markdown("""
                        <div style="background-color:#f0fff4; padding:10px; border-radius:5px; border:1px solid #c3e6cb;">
                        <p style="font-weight:bold; color:#2ecc71;">Corrected:</p>
                        </div>
                        """, unsafe_allow_html=True)
                        st.markdown(f"<div style='background-color:#fff; padding:10px; border-radius:5px; border:1px solid #eee;'>{isoChange['Corrected']}</div>", unsafe_allow_html=True)
                    if 'Reasoning' in isoChange and isoChange['Reasoning']:
                        st.markdown(f"""
                        <div style="background-color:#F0F8FF; padding:10px; border-radius:5px; margin-top:10px; border:1px solid #B0C4DE;">
                        <p style="font-weight:bold; color:#4169E1;">Error Type: {isoChange['Reasoning']}</p>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.info("No grammar errors were found in the document.")
                
    def _GenerateMarkdownReport(self, isoResults):
        isoMarkdown = f"# Grammar Correction Report\n\n"
        isoMarkdown += f"## Summary\n\n"
        isoTotalCorrections = len(isoResults['analysis']['changes'])
        isoBaseFilename = st.session_state.filename
        if '.' in isoBaseFilename:
            isoBaseFilename = isoBaseFilename.split('.')[0]
        isoMarkdown += f"* **Document Name**: {st.session_state.filename}\n"
        isoMarkdown += f"* **Total Corrections**: {isoTotalCorrections}\n"
        isoMarkdown += f"* **Analysis Date**: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        if isoResults['analysis']['error_types']:
            isoMarkdown += f"## Error Types Found\n\n"
            isoMarkdown += f"| Error Type | Count |\n"
            isoMarkdown += f"|------------|-------|\n"
            isoSortedErrorTypes = sorted(
                isoResults['analysis']['error_types'].items(),
                key=lambda x: x[1],
                reverse=True
            )
            for isoErrorType, isoCount in isoSortedErrorTypes:
                isoMarkdown += f"| {isoErrorType} | {isoCount} |\n"
        if isoResults['analysis']['changes']:
            isoMarkdown += f"\n## Detailed Corrections\n\n"
            for i, isoChange in enumerate(isoResults['analysis']['changes']):
                if not isoChange.get('Original', '') and not isoChange.get('Corrected', ''):
                    continue
                isoMarkdown += f"### Paragraph {isoChange.get('paragraph', i+1)}\n\n"
                isoMarkdown += f"**Original:**\n```\n{isoChange.get('Original', '')}\n```\n\n"
                isoMarkdown += f"**Corrected:**\n```\n{isoChange.get('Corrected', '')}\n```\n\n"
                if 'Reasoning' in isoChange and isoChange['Reasoning']:
                    isoMarkdown += f"**Error Type:** {isoChange['Reasoning']}\n\n"
                elif 'DetectedErrorTypes' in isoChange and isoChange['DetectedErrorTypes']:
                    primaryType = isoChange['DetectedErrorTypes'][0]
                    isoMarkdown += f"**Error Type:** {primaryType}\n\n"
                isoMarkdown += "---\n\n"
        return isoMarkdown
    
    def Run(self):
        isoConfig = self.RenderSidebar()
        if st.session_state.current_state == "upload":
            self.RenderUploadPage(isoConfig)
        elif st.session_state.current_state == "processing":
            self.RenderProcessingPage(isoConfig)
        elif st.session_state.current_state == "results":
            self.RenderResultsPage(isoConfig)

def CreateStreamlitApp():
    isoApp = DocumentCorrectionAppView()
    isoApp.Run()
    
if __name__ == "__main__":
    CreateStreamlitApp()